from pydantic import BaseModel
from typing import List, Dict, Any

class ParsedPage(BaseModel):
    page_number: int
    text: str
    sections: List[str]

class ParsedDocument(BaseModel):
    pages: List[ParsedPage]
    total_pages: int
    metadata: Dict[str, Any]

def _extract_sections(text: str) -> List[str]:
    # Look for all-caps lines or known headings
    sections = []
    known = ['RISK FACTORS', 'MANAGEMENT DISCUSSION AND ANALYSIS', 'FINANCIAL STATEMENTS', 'REVENUE', 'COMPETITION']
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.isupper() and len(line) > 3:
            sections.append(line)
        elif any(k in line.upper() for k in known):
            sections.append(line)
    return sections

def parse_pdf(file_path: str) -> ParsedDocument:
    """Parse PDF with page-level extraction. Filters empty pages."""
    try:
        import pymupdf as fitz
    except ImportError:
        import fitz
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text()
        # Clean excessive whitespace but preserve meaningful text
        text = ' '.join(text.split())
        # Skip empty or near-empty pages (likely blank or image-only)
        if len(text.strip()) < 10:
            continue
        pages.append(ParsedPage(
            page_number=i+1,
            text=text,
            sections=_extract_sections(text)
        ))
    total_pages = len(doc)
    doc.close()
    return ParsedDocument(pages=pages, total_pages=total_pages, metadata={'type': 'pdf'})

def parse_docx(file_path: str) -> ParsedDocument:
    """Parse DOCX with paragraph-level extraction."""
    import docx
    doc = docx.Document(file_path)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    full_text = "\n".join(paragraphs)
    full_text = ' '.join(full_text.split())
    
    pages = []
    chunk_size = 3000
    for i in range(0, len(full_text), chunk_size):
        chunk = full_text[i:i+chunk_size]
        if len(chunk.strip()) < 10:
            continue
        pages.append(ParsedPage(
            page_number=(i//chunk_size)+1,
            text=chunk,
            sections=_extract_sections(chunk)
        ))
    return ParsedDocument(pages=pages or [ParsedPage(page_number=1, text="", sections=[])], total_pages=max(1, len(pages)), metadata={'type': 'docx'})

def parse_txt(file_path: str) -> ParsedDocument:
    """Parse plain text files with encoding fallback."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            text = f.read()
    
    text = ' '.join(text.split())
    
    pages = []
    chunk_size = 3000
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i+chunk_size]
        if len(chunk.strip()) < 10:
            continue
        pages.append(ParsedPage(
            page_number=(i//chunk_size)+1,
            text=chunk,
            sections=_extract_sections(chunk)
        ))
    return ParsedDocument(pages=pages or [ParsedPage(page_number=1, text="", sections=[])], total_pages=max(1, len(pages)), metadata={'type': 'txt'})

def parse_document(file_path: str) -> ParsedDocument:
    ext = file_path.split('.')[-1].lower()
    if ext == 'pdf':
        return parse_pdf(file_path)
    elif ext == 'docx':
        return parse_docx(file_path)
    elif ext == 'txt':
        return parse_txt(file_path)
    raise ValueError("Unsupported file type")
